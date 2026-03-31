"""
app/rag/chunker.py
------------------
Document chunking — chia tài liệu thành các đoạn nhỏ (chunks) để index.
Hỗ trợ PDF, HTML, TXT, Markdown.
"""

from __future__ import annotations

import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import os
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
import io
import PyPDF2
from bs4 import BeautifulSoup

from app.core.config import get_settings
from app.core.logging import get_logger

logger = get_logger(__name__)
settings = get_settings()


@dataclass
class Chunk:
    """Đơn vị văn bản sau khi chunk.

    Attributes:
        id: ID duy nhất của chunk (UUID v4).
        text: Nội dung văn bản của chunk.
        metadata: Thông tin bổ sung (tên file, trang, thứ tự chunk, ...).
        token_count: Số token ước tính (dùng cho trim context).
    """

    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    token_count: int = 0

    def __post_init__(self) -> None:
        """Tính token_count nếu chưa được set."""
        if self.token_count == 0:
            # Ước tính thô: 1 token ≈ 4 ký tự
            self.token_count = len(self.text) // 4


class DocumentChunker:
    """Chia tài liệu thành Chunk với overlap.

    Sử dụng LangChain RecursiveCharacterTextSplitter bên dưới,
    nhưng wrap lại để trả về list[Chunk] với metadata thống nhất.

    Attributes:
        chunk_size (int): Số ký tự tối đa mỗi chunk.
        chunk_overlap (int): Số ký tự overlap giữa các chunk liên tiếp.
    """

    def __init__(
        self,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
    ) -> None:
        """Khởi tạo chunker với tuỳ chỉnh kích thước.

        Args:
            chunk_size: Kích thước chunk. Mặc định từ config (``DEFAULT_CHUNK_SIZE``).
            chunk_overlap: Overlap giữa chunks. Mặc định từ config (``DEFAULT_CHUNK_OVERLAP``).
        """
        self.chunk_size = chunk_size or settings.default_chunk_size
        self.chunk_overlap = chunk_overlap or settings.default_chunk_overlap

    def _load_file(self, file_path: Path) -> str:
        """Load và extract text từ file theo extension.

        Hỗ trợ: ``.pdf``, ``.html``, ``.htm``, ``.txt``, ``.md``.

        Args:
            file_path: Đường dẫn tới file.

        Returns:
            Văn bản raw đã được extract.

        Raises:
            ValueError: Nếu extension không được hỗ trợ.
            FileNotFoundError: Nếu file không tồn tại.
        """
        if not file_path.exists() or not file_path.is_file():
            raise FileNotFoundError(f"Không tìm thấy file tại đường dẫn: {file_path}")

        # Lấy đuôi file (ví dụ: '.pdf', '.txt') và chuyển về chữ thường để dễ so sánh
        ext = file_path.suffix.lower()
        data = ""

        if ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                data = f.read().strip()
                
        elif ext == ".pdf":
            reader = PyPDF2.PdfReader(str(file_path))
            for page in reader.pages:
                page_text = page.extract_text().strip()
                if page_text:
                    data += page_text.strip() + "\n"
                    
        elif ext == ".docx":
            doc = DocxDocument(str(file_path))
            for para in doc.paragraphs:
                if para.text.strip():
                    data += para.text.strip() + "\n"
                    
        elif ext in [".html", ".htm"]:
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()
                # Dùng BeautifulSoup lột bỏ các thẻ <div>, <p>, chỉ lấy chữ
                soup = BeautifulSoup(html_content, "html.parser")
                data = soup.get_text(separator="\n", strip=True)
                
        else:
            # lỗi ValueError nếu đuôi file không lọt vào các trường hợp trên
            raise ValueError(
                f"Định dạng file '{ext}' không được hỗ trợ. "
                f"Chỉ hỗ trợ: .pdf, .docx, .html, .htm, .txt, .md"
            )
                
        # Xóa các khoảng trắng và dấu xuống dòng thừa ở hai đầu văn bản
        return data.strip()
            

    def _split_text(self, text: str) -> list[str]:
        """Chia text thành list string theo chunk_size và chunk_overlap.

        Args:
            text: Toàn bộ văn bản cần chia.

        Returns:
            List string, mỗi phần tử là nội dung một chunk.
        """
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = self.chunk_size,
            chunk_overlap = self.chunk_overlap,
            length_function = len,
            # Danh sách các ký tự dùng để cắt, ưu tiên từ trái sang phải
            separators=["\n\n", "\n", " ", ""] 
        )

        chunks = text_splitter.split_text(text)
        return chunks

    def chunk_text(
        self,
        text: str,
        base_metadata: dict[str, Any] | None = None,
    ) -> list[Chunk]:
        """Chunk một đoạn text thô thành list Chunk.

        Args:
            text: Văn bản cần chunk.
            base_metadata: Metadata base sẽ được merge vào mỗi chunk.

        Returns:
            List[Chunk] đã được gán ID và metadata.
        """
        base_metadata = base_metadata or {}
        raw_texts = self._split_text(text) 
        
        chunks: list[Chunk] = [] # đóng gói các đoạn text --> Chunk
        total = len(raw_texts)
        
        for i, raw_text in enumerate(raw_texts):
            # Tạo bản sao metadata và thêm thông tin thứ tự của chunk này
            chunk_meta = base_metadata.copy()
            chunk_meta["chunk_index"] = i
            chunk_meta["total_chunks"] = total
            
            # Khởi tạo đối tượng Chunk (id và token_count sẽ tự động được tạo)
            new_chunk = Chunk(
                text=raw_text,
                metadata=chunk_meta
            )
            chunks.append(new_chunk)
            
        return chunks

    def chunk_file(
        self,
        file_path: str | Path,
        extra_metadata: dict[str, Any] | None = None,
    ) -> list[Chunk]:
        """Load file và chunk thành list Chunk.

        Args:
            file_path: Đường dẫn tới file (PDF, HTML, TXT, MD).
            extra_metadata: Metadata bổ sung sẽ được merge vào mỗi chunk
                            (ví dụ: ``{"source": "legal_code", "year": 2024}``).

        Returns:
            List[Chunk] với metadata gồm ``filename``, ``file_type``,
            ``chunk_index``, ``total_chunks`` và các extra_metadata.
        """
        # 1. Chuẩn hóa đường dẫn thành đối tượng Path để dễ thao tác
        path_obj = Path(file_path)
        
        # 2. Tạo metadata cơ bản từ chính cái file đó
        base_metadata = {
            "filename": path_obj.name,               
            "file_type": path_obj.suffix.lower(),
        }
        
        # Gộp thêm metadata từ bên ngoài truyền vào (nếu có)
        if extra_metadata:
            base_metadata.update(extra_metadata)
            
        # 3. Đọc chữ từ file (Hàm này có thể báo lỗi nếu file không tồn tại)
        raw_text = self._load_file(path_obj)
        
        # 4. Giao chữ và metadata cho hàm chunk_text xử lý và trả về kết quả
        chunks = self.chunk_text(
            text=raw_text, 
            base_metadata=base_metadata
        )
        
        return chunks

    def chunk_bytes(
        self,
        content: bytes,
        filename: str,
        extra_metadata: dict[str, Any] | None = None,
    ) -> list[Chunk]:
        """Chunk từ bytes (dùng khi nhận file upload qua API).

        Args:
            content: Nội dung file dạng bytes.
            filename: Tên file gốc (dùng để detect extension).
            extra_metadata: Metadata bổ sung.

        Returns:
            List[Chunk].
        """
        suffix = Path(filename).suffix.lower()
        base_metadata = {
            "filename": filename,
            "file_type": suffix,
        }
        if extra_metadata:
            base_metadata.update(extra_metadata)
            
        #Đọc chữ trực tiếp từ RAM (Bytes)
        raw_text = ""
        
        if suffix in [".txt", ".md"]:
            raw_text = content.decode("utf-8")
            
        #Đọc chữ trực tiếp từ RAM (Bytes)
        elif suffix == ".pdf":
            # Dùng io.BytesIO để biến cục bytes thành một "file ảo" trên RAM
            # Nhờ vậy thư viện đọc PDF vẫn tưởng nó đang đọc một file thật
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            pages_text = [page.extract_text().strip() for page in reader.pages if page.extract_text()]
            raw_text = "\n".join(pages_text)
            
        elif suffix in [".html", ".htm"]:
            # Giải mã HTML và dùng BeautifulSoup để lột bỏ hết các thẻ <div>, <p>... chỉ lấy chữ
            html_str = content.decode("utf-8")
            soup = BeautifulSoup(html_str, "html.parser")
            raw_text = soup.get_text(separator="\n", strip=True)
            
        elif suffix == ".docx":
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)
            raw_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
                
        else:
            # Chặn đứng nếu người dùng upload file không hỗ trợ (VD: .exe, .png)
            raise ValueError(f"Hệ thống chưa hỗ trợ xử lý định dạng file: {suffix}")
            
        # 4. Giao toàn bộ chữ thô và metadata cho "thợ đóng gói" xử lý
        chunks = self.chunk_text(
            text=raw_text,
            base_metadata=base_metadata
        )
        
        return chunks


if __name__ == "__main__":
    import pprint
    documentchunker = DocumentChunker()
    pprint.pprint(documentchunker.chunk_file(file_path = "D:\\TYP-Legal-Chatbot\\backend\\app\\rag\\data\\luat_lao_dong.txt")[0])